import os
import re
import uuid
import shutil
import zipfile
import subprocess # 🚀 改用這個來呼叫 Docker 裡的 LibreOffice
from flask import Flask, request, send_file, after_this_request
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

# 🚀 Cloud Run 環境建議使用 /tmp，這是 Linux 容器內可寫入的記憶體空間
UPLOAD_FOLDER = '/tmp/temp_uploads'

def init_storage():
    if os.path.exists(UPLOAD_FOLDER):
        try:
            shutil.rmtree(UPLOAD_FOLDER)
        except Exception as e:
            print(f"清理失敗: {e}")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_number(filename):
    match = re.search(r'附件(\d+)', filename)
    return int(match.group(1)) if match else None

# 🚀 關鍵修正：將 win32 轉檔換成 Linux 指令
def convert_doc_to_docx(folder_path):
    files = [f for f in os.listdir(folder_path) if f.lower().endswith('.doc') and not f.startswith('~$')]
    for f in files:
        abs_path = os.path.abspath(os.path.join(folder_path, f))
        try:
            # 呼叫 Dockerfile 裡安裝的 libreoffice 進行轉檔
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx', 
                abs_path, '--outdir', folder_path
            ], check=True)
            print(f"Linux 轉檔完成: {f}")
        except Exception as e:
            print(f"Linux 轉檔出錯: {e}")

@app.route('/process', methods=['POST'])
def process():
    session_id = str(uuid.uuid4())
    user_dir = os.path.join(UPLOAD_FOLDER, session_id)
    os.makedirs(user_dir, exist_ok=True)

    uploaded_files = request.files.getlist('files')
    for file in uploaded_files:
        if file.filename:
            safe_filename = os.path.basename(file.filename)
            file.save(os.path.join(user_dir, safe_filename))

    # 執行 Linux 轉檔
    convert_doc_to_docx(user_dir)

    docx_files = [f for f in os.listdir(user_dir) if f.endswith('.docx') and not f.startswith('~$')]
    for filename in docx_files:
        num = extract_number(filename)
        if num is not None:
            try:
                doc = Document(os.path.join(user_dir, filename))
                for section in doc.sections:
                    header = section.header
                    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    p.text = f"附件{num}"
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.save(os.path.join(user_dir, filename))
            except Exception as e:
                print(f"標註 {filename} 失敗: {e}")

    zip_filename = f"{session_id}.zip"
    zip_path = os.path.join(UPLOAD_FOLDER, zip_filename)
    
    with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as z:
        for f in os.listdir(user_dir):
            if f.endswith('.docx'):
                z.write(os.path.join(user_dir, f), f)

    @after_this_request
    def cleanup(response):
        try:
            shutil.rmtree(user_dir)
            if os.path.exists(zip_path):
                os.remove(zip_path)
        except: pass
        return response

    return send_file(zip_path, as_attachment=True, download_name="processed_files.zip")

@app.route('/convert-pdf', methods=['POST'])
def convert_pdf():
    session_id = str(uuid.uuid4())
    user_dir = os.path.join(UPLOAD_FOLDER, session_id)
    os.makedirs(user_dir, exist_ok=True)

    files = request.files.getlist('files')
    for file in files:
        if file.filename:
            safe_name = os.path.basename(file.filename)
            file.save(os.path.join(user_dir, safe_name))

    # 🚀 使用 LibreOffice 指令直接將所有 doc/docx 轉為 pdf
    try:
        # --convert-to pdf 會自動處理資料夾內支援的格式
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', 
            os.path.join(user_dir, '*'), '--outdir', user_dir
        ], shell=True, check=True)
    except Exception as e:
        print(f"PDF 轉檔失敗: {e}")

    zip_path = os.path.join(UPLOAD_FOLDER, f"{session_id}_pdf.zip")
    with zipfile.ZipFile(zip_path, 'w') as z:
        for f in os.listdir(user_dir):
            if f.lower().endswith('.pdf'):
                z.write(os.path.join(user_dir, f), f)

    @after_this_request
    def cleanup(response):
        shutil.rmtree(user_dir, ignore_errors=True)
        if os.path.exists(zip_path): os.remove(zip_path)
        return response

    return send_file(zip_path, as_attachment=True, download_name="converted_pdfs.zip")


if __name__ == '__main__':
    init_storage()
    # 🚀 確保讀取 Cloud Run 的 PORT
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)